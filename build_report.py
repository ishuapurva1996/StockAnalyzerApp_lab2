"""Generate Lab2_Report.docx with screenshot placeholders.

Run:
    /Library/Frameworks/Python.framework/Versions/Current/bin/python3 build_report.py
"""

import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


# ---------- helpers ----------------------------------------------------------

def _set_cell_border(cell, **kwargs):
    """Set cell borders via low-level XML."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = tc_pr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tc_pr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"), "single")
            tag.set(qn("w:sz"), "8")
            tag.set(qn("w:color"), "888888")
            existing = tcBorders.find(qn(f"w:{edge}"))
            if existing is not None:
                tcBorders.remove(existing)
            tcBorders.append(tag)


def _shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def screenshot_box(doc, label, height_cm=8.0):
    """Insert a single-cell bordered table with placeholder text the user can
    click to delete, then paste a screenshot in its place."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(16.0)
    cell = table.rows[0].cells[0]
    cell.width = Cm(16.0)
    cell.height = Cm(height_cm)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_border(cell, top=True, left=True, bottom=True, right=True)
    _shade_cell(cell, "F8F8F8")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ Insert screenshot: {label} ]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.size = Pt(11)

    # Force minimum height
    tr = table.rows[0]._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_cm * 567)))  # cm to twips
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Menlo"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def section(doc, heading, level=1):
    h = doc.add_heading(heading, level=level)
    return h


def feature_block(doc, title, description, file_ref, screenshot_label,
                  caption, comment, height_cm=8.0):
    section(doc, title, level=2)
    doc.add_paragraph(description)
    if file_ref:
        p = doc.add_paragraph()
        p.add_run("File: ").bold = True
        run = p.add_run(file_ref)
        run.font.name = "Menlo"
        run.font.size = Pt(10)
    screenshot_box(doc, screenshot_label, height_cm=height_cm)
    add_caption(doc, caption)
    if comment:
        p = doc.add_paragraph()
        p.add_run("Comment: ").bold = True
        p.add_run(comment)
    doc.add_paragraph()  # spacer


def test_block(doc, phase_num, title, what, how_to_run, screenshot_label,
               caption, height_cm=9.0):
    section(doc, f"Phase {phase_num} — {title}", level=2)
    p = doc.add_paragraph()
    p.add_run("What is tested: ").bold = True
    p.add_run(what)
    p = doc.add_paragraph()
    p.add_run("How to run: ").bold = True
    add_code(doc, how_to_run)
    screenshot_box(doc, screenshot_label, height_cm=height_cm)
    add_caption(doc,
                caption + "  (Note: the last line of the test output prints "
                          "'Execution time: X.XXXX seconds (X.XX ms)' — make "
                          "sure that line is visible in the screenshot.)")
    doc.add_paragraph()


# ---------- build the document ----------------------------------------------

def build():
    doc = Document()

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ============== Cover ==============
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("\n\n\nDATA 200")
    run.font.size = Pt(20)
    run.bold = True
    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title2.add_run("San Jose State University")
    run.font.size = Pt(14)

    title3 = doc.add_paragraph()
    title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title3.add_run("\nLab 2 — Stock Analysis")
    run.font.size = Pt(28)
    run.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Web Scraping & CSV Import\nImplementation, Sample Output, and Unit-Test Results")
    run.font.size = Pt(13)
    run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Student: ____________________________\n").font.size = Pt(12)
    info.add_run("Date:    ____________________________\n").font.size = Pt(12)
    info.add_run("Course:  DATA 200").font.size = Pt(12)

    doc.add_page_break()

    # ============== Project Overview ==============
    section(doc, "Project Overview", level=1)
    doc.add_paragraph(
        "Lab 2 extends the Stock Analysis project from manual entry to two "
        "automated data-loading paths: web scraping from Yahoo! Finance using "
        "Selenium + BeautifulSoup, and importing a Yahoo!-format CSV file. "
        "Both paths are exposed in the GUI (tkinter Web menu) and the console "
        "(Manage Data sub-menu). The implementation spans seven build phases, "
        "each with its own automated test suite. This document captures one "
        "screenshot per requested functionality plus one screenshot per phase "
        "of unit-test output, with execution time clearly visible."
    )
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("File map:").bold = True
    files = [
        ("stock_class.py", "Stock and DailyData classes; built-in unit tests"),
        ("utilities.py",   "sortStocks, sortDailyData, display_stock_chart, clear_screen"),
        ("stock_data.py",  "retrieve_stock_web (web scrape), import_stock_web_csv (CSV), SQLite save/load"),
        ("stock_GUI.py",   "tkinter GUI: menubar, listbox, notebook, buy/sell/delete, scrape/import/chart"),
        ("stock_console.py","Console UI: menus, add/buy/sell/delete/list, retrieve_from_web, import_csv"),
        ("stocks.py",      "Single entry point that launches GUI or console"),
    ]
    for f, desc in files:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f)
        run.font.name = "Menlo"
        run.bold = True
        p.add_run(f" — {desc}")

    doc.add_page_break()

    # ============== Class diagram (IS-A / HAS-A) ==============
    section(doc, "Class Diagram — IS-A and HAS-A Relationships", level=1)
    doc.add_paragraph(
        "The diagram below shows the three classes that make up the Stock "
        "Analyzer along with their inheritance (IS-A) and composition (HAS-A) "
        "relationships. An editable copy is provided alongside this report as "
        "class_diagram.drawio (open at https://app.diagrams.net/)."
    )

    # Embed the PNG if present
    diagram_path = "class_diagram.png"
    if os.path.exists(diagram_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(diagram_path, width=Inches(6.5))
        add_caption(doc,
                    "Figure 0.1 — Stock, DailyData, and StockApp. Hollow "
                    "triangles denote IS-A (inheritance); hollow diamonds "
                    "denote HAS-A (composition).")
    else:
        screenshot_box(doc, "Insert class_diagram.png here", height_cm=10.0)
        add_caption(doc, "Figure 0.1 — Class diagram.")

    # IS-A explanation
    p = doc.add_paragraph()
    p.add_run("IS-A relationships (inheritance):").bold = True
    items = [
        ("Stock IS-A object", "Stock is defined as a regular Python class; in "
                              "Python 3 every class implicitly inherits from "
                              "object, so this gives Stock its default "
                              "lifecycle methods (__init__, __repr__, …)."),
        ("DailyData IS-A object", "Same as Stock — inherits the default "
                                  "Python object behaviour."),
        ("StockApp IS-A object", "The GUI controller class also inherits "
                                 "implicitly from object. (It does not "
                                 "subclass tkinter.Tk; it composes a Tk "
                                 "instance instead — see HAS-A below.)"),
    ]
    for title, desc in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(title)
        run.bold = True
        p.add_run(" — " + desc)

    # HAS-A explanation
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("HAS-A relationships (composition):").bold = True
    items = [
        ("Stock HAS-A list of DailyData",
         "Each Stock owns its own price history. The container is the "
         "Stock.DataList attribute (initialized as [] in __init__) and is "
         "extended via the public method add_data(stock_data: DailyData). "
         "Cardinality is 0..* — a freshly added stock has no daily rows."),
        ("StockApp HAS-A list of Stock",
         "The GUI’s portfolio is held in StockApp.stock_list. add_stock() "
         "appends a new Stock; delete_stock() removes one; load() fills the "
         "list from the SQLite database. The same list is passed by "
         "reference into the data layer functions (save_stock_data, "
         "load_stock_data, retrieve_stock_web, import_stock_web_csv)."),
        ("StockApp HAS-A several tkinter widgets",
         "The GUI controller composes (does not subclass) Tk plus its menus, "
         "Listbox, ttk.Notebook, and Text/Entry widgets. This is a deliberate "
         "design choice: composition keeps the controller decoupled from "
         "tkinter’s class hierarchy and makes the widgets easy to test "
         "independently."),
    ]
    for title, desc in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(title)
        run.bold = True
        p.add_run(" — " + desc)

    # Why the design works
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Why this matters: ").bold = True
    p.add_run(
        "Encapsulating daily history inside Stock (HAS-A) means every "
        "function that consumes a Stock — chart rendering, report formatting, "
        "save/load, web scraping, CSV import — works against the same shape "
        "of data. Likewise, holding the portfolio inside StockApp via a "
        "passed-by-reference list lets the data-layer functions populate "
        "the same in-memory structure that the GUI is rendering, with no "
        "extra copying or marshalling."
    )

    doc.add_page_break()

    # ============== SECTION 1: Web Scraping ==============
    section(doc, "1. Web Scraping (Yahoo! Finance)", level=1)
    doc.add_paragraph(
        "The web-scraping pipeline launches Chrome via Selenium, retrieves "
        "the Yahoo! Finance history page, and parses its HTML table with "
        "BeautifulSoup. Modern Selenium (≥ 4.10) ships with Selenium Manager, "
        "so the matching ChromeDriver is downloaded automatically — no manual "
        "PATH configuration is needed."
    )

    feature_block(
        doc,
        "1.1  retrieve_stock_web() — the scraper",
        "Builds the Yahoo! Finance history URL from epoch-second start/end "
        "dates, launches headless-friendly Chrome with logging and JavaScript "
        "preferences disabled, then walks every <tr> in the result table. "
        "Rows containing exactly seven <td> cells are real price rows "
        "(dividends/splits have a different shape and are skipped). For each "
        "row a DailyData(date, close, volume) is created and added to the "
        "matching Stock.",
        "stock_data.py  (lines 92–122)",
        "Live MSFT scrape via GUI Web → Scrape Data from Yahoo! Finance",
        "Figure 1.1 — Scrape Data from Yahoo! Finance with date range and result.",
        "After entering the start and end dates, Chrome opens, the history "
        "page loads, and the History/Report tabs show the new daily rows. A "
        "“Data Retrieved” confirmation is shown when finished."
    )

    feature_block(
        doc,
        "1.2  scrape_web_data() — GUI wiring",
        "Hooked into the Web menu of stock_GUI.py. Uses simpledialog.askstring "
        "to prompt for the start and end dates in m/d/yy format, calls "
        "retrieve_stock_web(), then calls display_stock_data() to refresh the "
        "History and Report tabs. If the Chrome driver is missing, an error "
        "dialog explains the cause without crashing the app.",
        "stock_GUI.py  (lines 141–150)",
        "GUI: Web menu open with both options visible",
        "Figure 1.2 — Web menu on the main window.",
        "The same retrieve_stock_web() function powers both this GUI flow "
        "and the console flow in section 1.3."
    )

    feature_block(
        doc,
        "1.3  retrieve_from_web() — console wiring",
        "The Manage Data → Retrieve Data from Web option in stock_console.py. "
        "Prompts for MM/DD/YY start and end dates, calls "
        "stock_data.retrieve_stock_web(), then prints the number of records "
        "retrieved. Catches RuntimeWarning so a missing Chrome driver "
        "produces a friendly message instead of a stack trace.",
        "stock_console.py  (retrieve_from_web)",
        "Console: Manage Data → Retrieve Data from Web with output",
        "Figure 1.3 — “Records Retrieved: N” printed after the scrape completes.",
        "The console flow mirrors the PDF screenshot exactly: prompts for "
        "start and end dates, then prints the record count."
    )

    doc.add_page_break()

    # ============== SECTION 2: CSV Import ==============
    section(doc, "2. CSV File Import", level=1)
    doc.add_paragraph(
        "Yahoo! Finance also lets you download price history as a CSV file. "
        "The CSV-import path is faster, more reliable, and does not require "
        "Chrome — useful when the user has already downloaded the file from "
        "the Yahoo! Finance “Historical Data” tab."
    )

    feature_block(
        doc,
        "2.1  import_stock_web_csv() — the parser",
        "Opens the CSV with a `with` block, creates a csv.reader using comma "
        "as the delimiter, calls next() once to skip the header row, then for "
        "each remaining row constructs DailyData(date=row[0], close=row[4], "
        "volume=row[6]) and appends it to the matching Stock. The expected "
        "Yahoo! schema is Date, Open, High, Low, Close, Adj Close, Volume.",
        "stock_data.py  (lines 125–133)",
        "Sample AAPL.csv preview (Date, Open, High, Low, Close, Adj Close, Volume)",
        "Figure 2.1 — Yahoo! Finance CSV opened in a viewer for reference.",
        "The closing price is column index 4 (the 5th column) and volume is "
        "column index 6 (the 7th column); column index 5 (Adj Close) is "
        "deliberately ignored."
    )

    feature_block(
        doc,
        "2.2  importCSV_web_data() — GUI wiring",
        "Bound to the Web → Import CSV From Yahoo! Finance menu item. Reads "
        "the symbol from the currently selected listbox row, opens "
        "filedialog.askopenfilename to choose the .csv file, calls "
        "import_stock_web_csv(), refreshes the History/Report tabs, and shows "
        "an “Import Complete” confirmation.",
        "stock_GUI.py  (lines 153–159)",
        "GUI: file picker open, AAPL.csv selected, Import Complete dialog",
        "Figure 2.2 — File picker filtered to *.csv plus the success dialog.",
        "Selecting the stock in the listbox before opening the picker is "
        "important — the function uses the listbox selection to decide which "
        "Stock to attach data to."
    )

    feature_block(
        doc,
        "2.3  import_csv() — console wiring",
        "Manage Data → Import from CSV File. Prints the current stock list, "
        "prompts for the symbol to use, prompts for the file path, then "
        "calls import_stock_web_csv() and prints “CSV File Imported”.",
        "stock_console.py  (import_csv)",
        "Console: Import CSV interaction with “CSV File Imported” line",
        "Figure 2.3 — Symbol and filename prompts plus the import confirmation.",
        "If the path is wrong, the function reports a clean error message "
        "(file not found / OS error) rather than crashing."
    )

    doc.add_page_break()

    # ============== SECTION 3: Console sample outputs ==============
    section(doc, "3. Console Sample Outputs (PDF §2.6)", level=1)
    doc.add_paragraph(
        "These screenshots reproduce the sample outputs from section 2.6 of "
        "the lab handout, demonstrating each console feature working end-to-end."
    )

    feature_block(
        doc, "3.1  Main menu (Stock Analyzer)",
        "Top-level menu listing the five workflows plus exit.",
        "stock_console.py  (main_menu)",
        "Console: Stock Analyzer main menu",
        "Figure 3.1 — Stock Analyzer main menu.",
        "Options 1–5 are dispatched to manage_stocks, add_stock_data, "
        "display_report, display_chart, and manage_data respectively."
    )
    feature_block(
        doc, "3.2  Manage Stocks → Add Stock",
        "Option 1 → 1: prompt for symbol, name, shares; loop until 0 to stop.",
        "stock_console.py  (add_stock)",
        "Console: Add Stock — MSFT entered with 60 shares",
        "Figure 3.2 — Add Stock prompt after one entry.",
        "The symbol is uppercased automatically; non-numeric share input is "
        "rejected with a retry."
    )
    feature_block(
        doc, "3.3  Manage Stocks → Update Shares (Buy / Sell)",
        "Sub-menu with 1 = Buy, 2 = Sell. Both prompts show the current stock "
        "list inside square brackets, then ask for the symbol and amount.",
        "stock_console.py  (update_shares, buy_stock, sell_stock)",
        "Console: Update Shares menu and Buy Shares interaction",
        "Figure 3.3 — Buy Shares for MSFT, 40 added.",
        "Buy increases self._shares; Sell decreases it. Both go through the "
        "Stock class’s buy()/sell() methods, which are the only sanctioned "
        "way to change the share count."
    )
    feature_block(
        doc, "3.4  Manage Stocks → Delete Stock",
        "Prints the stock list, prompts for the symbol, and removes it.",
        "stock_console.py  (delete_stock)",
        "Console: Delete Stock prompt",
        "Figure 3.4 — Delete Stock with stock list shown.",
        "Unknown symbols print a “Stock not found” message rather than "
        "deleting silently."
    )
    feature_block(
        doc, "3.5  Manage Stocks → List Stocks",
        "Tabular dump of SYMBOL, NAME, SHARES across all stocks.",
        "stock_console.py  (list_stocks)",
        "Console: Stock List — header and rows",
        "Figure 3.5 — List Stocks showing MSFT MICROSOFT 100.0.",
        "Header columns are aligned with str.format width specifiers."
    )
    feature_block(
        doc, "3.6  Add Daily Stock Data",
        "Prompts for the symbol, then loops on “Enter Date,Price,Volume:” "
        "until a blank line is entered. Each line is parsed into a DailyData "
        "and added to the Stock.",
        "stock_console.py  (add_stock_data)",
        "Console: Add Daily Stock Data prompt with example",
        "Figure 3.6 — Comma-separated entry prompt.",
        "Date format is m/d/yy (e.g. 8/28/20). Invalid lines print an error "
        "and re-prompt without losing earlier entries."
    )
    feature_block(
        doc, "3.7  Show Report",
        "Per-stock report: symbol + name, shares, daily rows, then "
        "start/end/high/low/profit-loss block.",
        "stock_console.py  (display_report)",
        "Console: Stock Report for MSFT (or “*** No daily history.”)",
        "Figure 3.7 — Stock Report output.",
        "Profit/Loss is signed (e.g. -$700.00) and formatted with thousands "
        "separators."
    )
    feature_block(
        doc, "3.8  Manage Data (Save, Load, Retrieve, Import CSV)",
        "Five-option sub-menu wrapping the SQLite save/load and the two "
        "Lab-2 data-loading paths.",
        "stock_console.py  (manage_data)",
        "Console: Manage Data sub-menu and a Save → Load round-trip",
        "Figure 3.8 — Manage Data menu plus “--- Data Loaded from Database ---”.",
        "The save/load cycle uses sqlite3 with a composite primary key on "
        "(symbol, date) so importing the same window twice is harmless."
    )
    feature_block(
        doc, "3.9  Show Chart — Console",
        "From the main menu, option 4 prompts for a symbol and calls "
        "utilities.display_stock_chart(), which uses matplotlib to draw a "
        "line chart of close prices vs. date. The chart opens in its own "
        "window separate from the terminal.",
        "stock_console.py  (display_chart) → utilities.display_stock_chart",
        "Chart window: MICROSOFT line chart (Date / Price)",
        "Figure 3.9 — MSFT chart launched from the console menu, matching the "
        "PDF screenshot.",
        "The chart’s title is the company name; x-axis is Date, y-axis is "
        "Price.",
        height_cm=10.0,
    )

    feature_block(
        doc, "3.10  Show Chart — GUI",
        "The GUI exposes the same chart through the Chart menu → Display "
        "Stock Chart. The user selects a stock in the listbox first, then "
        "opens the menu; the same utilities.display_stock_chart() is called, "
        "so both interfaces produce identical figures from the same data.",
        "stock_GUI.py  (display_chart, lines 162–164) → "
        "utilities.display_stock_chart",
        "Chart window launched from GUI Chart menu — MICROSOFT line chart",
        "Figure 3.10 — Chart triggered from the GUI Chart → Display Stock "
        "Chart menu item.",
        "The line should plot Date on the x-axis and Price on the y-axis "
        "(see PDF page 17 reference). Ticks rotate to fit the date strings.",
        height_cm=10.0,
    )

    doc.add_page_break()

    # ============== SECTION 4: Unit Testing ==============
    section(doc, "4. Unit Testing — Per-Phase Results", level=1)
    doc.add_paragraph(
        "The implementation is covered by seven test files (test_phase1.py "
        "through test_phase7.py) totalling 125 assertions. Each test prints "
        "Successful! per check, a final “Congratulations — All Phase N Tests "
        "Passed” line, and an Execution time: X.XXXX seconds (X.XX ms) line. "
        "The execution-time line should be clearly visible in each unit-test "
        "screenshot below."
    )

    p = doc.add_paragraph()
    p.add_run("To run any single phase: ").bold = True
    add_code(doc,
             '/Library/Frameworks/Python.framework/Versions/Current/bin/python3 '
             'test_phase<N>.py')
    p = doc.add_paragraph()
    p.add_run("To run them all in sequence: ").bold = True
    add_code(doc,
             'for n in 1 2 3 4 5 6 7; do '
             '/Library/Frameworks/Python.framework/Versions/Current/bin/python3 '
             'test_phase$n.py; done')
    doc.add_paragraph()

    test_block(
        doc, 1, "Stock and DailyData classes",
        "The unit tests bundled inside stock_class.main(): construct a Stock, "
        "verify symbol/shares are immutable except via buy()/sell(), verify "
        "name is mutable, exercise buy/sell, and confirm DailyData fields "
        "round-trip correctly. 7 assertions.",
        '/Library/Frameworks/Python.framework/Versions/Current/bin/python3 test_phase1.py',
        "Phase 1 unit-test output with execution time",
        "Figure 4.1 — Phase 1 test run."
    )
    test_block(
        doc, 2, "GUI scaffolding",
        "Constructs StockApp with mainloop disabled and asserts that all "
        "widgets exist with the right Tk types: headingLabel, stockList, "
        "addSymbolEntry, addNameEntry, addSharesEntry, updateSharesEntry, "
        "dailyDataList, stockReport. Also verifies the menubar (File / Web / "
        "Chart cascades and their items) and the ttk.Notebook tabs (Main / "
        "History / Report) and the <<ListboxSelect>> binding. 19 assertions.",
        '/Library/Frameworks/Python.framework/Versions/Current/bin/python3 test_phase2.py',
        "Phase 2 unit-test output with execution time",
        "Figure 4.2 — Phase 2 test run."
    )
    test_block(
        doc, 3, "GUI completion (delete + report)",
        "Seeds two stocks (MSFT with 4 daily rows, AAPL with none) and "
        "asserts the report renders the right symbol, name, shares, "
        "start/end/high/low prices and signed profit-loss; also verifies "
        "“*** No daily history.” appears for the empty stock. Then patches "
        "messagebox.askyesno to test delete_stock for confirm=yes, "
        "confirm=no, and no-selection. 17 assertions.",
        '/Library/Frameworks/Python.framework/Versions/Current/bin/python3 test_phase3.py',
        "Phase 3 unit-test output with execution time",
        "Figure 4.3 — Phase 3 test run."
    )
    test_block(
        doc, 4, "Console stock management",
        "Patches input() and clear_screen() and asserts: add_stock appends a "
        "Stock with uppercase symbol; invalid shares are retried; buy_stock "
        "and sell_stock change the share count; unknown symbols are no-ops; "
        "delete_stock removes the right entry; list_stocks prints the right "
        "header and rows; the update_shares menu dispatches Buy and Sell "
        "correctly. 25 assertions.",
        '/Library/Frameworks/Python.framework/Versions/Current/bin/python3 test_phase4.py',
        "Phase 4 unit-test output with execution time",
        "Figure 4.4 — Phase 4 test run."
    )
    test_block(
        doc, 5, "Console data entry, report, chart",
        "add_stock_data parses comma-separated lines into DailyData "
        "(rejecting bad rows); display_report prints all expected substrings "
        "for both populated and empty stocks; display_chart calls "
        "display_stock_chart with the uppercased symbol and the full stock "
        "list. 23 assertions.",
        '/Library/Frameworks/Python.framework/Versions/Current/bin/python3 test_phase5.py',
        "Phase 5 unit-test output with execution time",
        "Figure 4.5 — Phase 5 test run."
    )
    test_block(
        doc, 6, "Console Lab-2 hooks",
        "Tests manage_data dispatch (Save, Load, invalid option), "
        "retrieve_from_web (date forwarding, record-count print, "
        "RuntimeWarning surfacing), and import_csv (symbol/filename "
        "forwarding plus confirmation). Also runs a real save → load "
        "round-trip against a SQLite file in a temp directory and verifies "
        "that the stock and its 2 daily rows survive intact. 20 assertions.",
        '/Library/Frameworks/Python.framework/Versions/Current/bin/python3 test_phase6.py',
        "Phase 6 unit-test output with execution time",
        "Figure 4.6 — Phase 6 test run."
    )
    test_block(
        doc, 7, "End-to-end (live web + chart + CSV)",
        "Three independent integration checks: (1) parse a synthetic "
        "Yahoo!-format CSV and verify all 5 daily rows landed; (2) render a "
        "matplotlib chart with the Agg backend (no GUI window) and assert on "
        "the figure’s title, axis labels, and data points; (3) launch "
        "headless Chrome via Selenium Manager, scrape MSFT history for "
        "12/01/24–12/15/24, and verify > 0 records were returned. 14 "
        "assertions.",
        '/Library/Frameworks/Python.framework/Versions/Current/bin/python3 test_phase7.py',
        "Phase 7 end-to-end test output with execution time",
        "Figure 4.7 — Phase 7 test run (note the “scraped N record(s)” line "
        "from the live Yahoo scrape).",
        height_cm=10.0,
    )

    # ============== Summary ==============
    doc.add_page_break()
    section(doc, "5. Summary", level=1)
    doc.add_paragraph(
        "All seven phases are implemented and verified by automated tests. "
        "125 assertions pass across the seven test files. Both Lab-2 "
        "deliverables — web scraping and CSV import — work in the GUI and "
        "console. The live Yahoo! Finance scrape was confirmed end-to-end on "
        "a recent date range with headless Chrome 147 and Selenium 4.43."
    )

    p = doc.add_paragraph()
    p.add_run("Per-phase assertion counts:").bold = True
    counts = [
        ("Phase 1 — Stock / DailyData unit tests", 7),
        ("Phase 2 — GUI scaffolding",              19),
        ("Phase 3 — GUI completion",               17),
        ("Phase 4 — Console stock management",     25),
        ("Phase 5 — Console data, report, chart",  23),
        ("Phase 6 — Console Lab-2 hooks",          20),
        ("Phase 7 — End-to-end",                   14),
    ]
    for label, n in counts:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(label + ": ")
        run = p.add_run(f"{n} assertions  ✓")
        run.bold = True
    p = doc.add_paragraph()
    run = p.add_run("Total: 125 assertions, all passing.")
    run.bold = True

    # ---- save ----
    out_path = "Lab2_Report.docx"
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    build()
